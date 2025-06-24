from docx import Document
from app.ai.summarizer import summarize_text
import os

def create_resume(data) -> dict:
    doc = Document()

    # === Header ===
    doc.add_heading(data.name, 0)
    doc.add_paragraph(data.title)

    contact_info = f"{data.phone} | {data.email} | {data.address}"
    doc.add_paragraph(contact_info)

    # === Links ===
    links = " | ".join(filter(None, [data.linkedin, data.github]))
    if links:
        doc.add_paragraph(links)

    # === Summary ===
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(summarize_text(data.summary))

    # === Experience ===
    doc.add_heading("Experience", level=1)
    for exp in data.experience:
        duration = f"{exp.from_month} {exp.from_year} to {exp.to_month} {exp.to_year}"
        doc.add_paragraph(f"{exp.job_title}, {exp.company} ({duration})", style="List Bullet")
        doc.add_paragraph(exp.details)

    # === Education ===
    doc.add_heading("Education", level=1)
    for edu in data.education:
        doc.add_paragraph(f"• {edu}")

    # === Technical Skills ===
    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph(", ".join(data.technical_skills))

    # === Soft Skills ===
    doc.add_heading("Soft Skills", level=1)
    doc.add_paragraph(", ".join(data.soft_skills))

    # === Certifications ===
    if data.certifications:
        doc.add_heading("Certifications", level=1)
        for cert in data.certifications:
            doc.add_paragraph(f"• {cert}")

    # === Projects ===
    if data.projects:
        doc.add_heading("Projects", level=1)
        for project in data.projects:
            doc.add_paragraph(f"• {project}")

    # === Languages ===
    if data.languages:
        doc.add_heading("Languages", level=1)
        doc.add_paragraph(", ".join(data.languages))

    # === Save Resume ===
    os.makedirs("static", exist_ok=True)
    safe_name = data.name.replace(" ", "_")
    filename = f"{safe_name}_resume.docx"
    filepath = os.path.join("static", filename)
    doc.save(filepath)

    return {
        "message": "Resume generated successfully",
        "filename": filename
    }
